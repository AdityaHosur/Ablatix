"""
Document Parser for Ablatix Indexer.
Implements layout-aware parsing with overlapping chunks for better context.
"""

import os
from typing import List, Dict, Any, Optional

try:
    import fitz  # PyMuPDF for PDF parsing
except ImportError:
    fitz = None

try:
    import docx  # python-docx for DOCX parsing
except ImportError:
    docx = None


class DocumentParser:
    def __init__(
        self,
        parser_type: str = "llama",
        chunk_size: int = 500,
        chunk_overlap: int = 100
    ):
        """
        Initialize the parser.
        :param parser_type: 'llama' or 'ragflow'
        :param chunk_size: Maximum number of characters per chunk.
        :param chunk_overlap: Number of overlapping characters between chunks.
        """
        self.parser_type = parser_type.lower()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        if self.parser_type not in ("llama", "ragflow"):
            raise ValueError("parser_type must be 'llama' or 'ragflow'")

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse a document and return overlapping chunks.
        :param file_path: Path to the document file.
        :return: List of parsed and chunked elements.
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            raw_blocks = self._parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            raw_blocks = self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

        # Apply overlapping chunking to parsed blocks
        return self._chunk_blocks(raw_blocks)

    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse a PDF file into raw blocks.
        """
        if fitz is None:
            raise ImportError("PyMuPDF (fitz) is required for PDF parsing.")
        doc = fitz.open(file_path)
        results = []
        for page_num, page in enumerate(doc, 1):
            blocks = page.get_text("blocks")
            for block in blocks:
                text = block[4].strip()
                if text:
                    results.append({
                        "type": "paragraph",
                        "text": text,
                        "metadata": {
                            "page": page_num,
                            "bbox": block[:4]
                        }
                    })
        doc.close()
        return results

    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse a DOCX file into raw blocks.
        """
        if docx is None:
            raise ImportError("python-docx is required for DOCX parsing.")
        doc = docx.Document(file_path)
        results = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                results.append({
                    "type": "paragraph",
                    "text": text,
                    "metadata": {
                        "style": para.style.name
                    }
                })
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            # Flatten table into text for embedding
            table_text = "\n".join([" | ".join(row) for row in table_data])
            if table_text.strip():
                results.append({
                    "type": "table",
                    "text": table_text,
                    "metadata": {
                        "table": table_data
                    }
                })
        return results

    def _chunk_blocks(self, blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Apply overlapping chunking to raw parsed blocks.
        Merges blocks into chunks of `chunk_size` characters,
        with `chunk_overlap` characters of overlap between chunks.
        :param blocks: List of raw parsed blocks.
        :return: List of overlapping chunks.
        """
        # Concatenate all block texts with their metadata
        full_text = ""
        char_to_meta = []  # Maps character index to block metadata

        for block in blocks:
            text = block.get("text", "")
            meta = block.get("metadata", {})
            block_type = block.get("type", "paragraph")
            start = len(full_text)
            full_text += text + "\n"
            end = len(full_text)
            char_to_meta.append({
                "start": start,
                "end": end,
                "metadata": meta,
                "type": block_type
            })

        # Generate overlapping chunks
        chunks = []
        start = 0
        total_length = len(full_text)

        while start < total_length:
            end = min(start + self.chunk_size, total_length)
            chunk_text = full_text[start:end].strip()

            if chunk_text:
                # Find the metadata for the block that starts this chunk
                meta = self._get_meta_for_position(char_to_meta, start)
                chunks.append({
                    "type": meta.get("type", "paragraph"),
                    "text": chunk_text,
                    "metadata": {
                        **meta.get("metadata", {}),
                        "chunk_start": start,
                        "chunk_end": end
                    }
                })

            # Move forward by chunk_size - chunk_overlap
            step = self.chunk_size - self.chunk_overlap
            if step <= 0:
                raise ValueError("chunk_size must be greater than chunk_overlap.")
            start += step

        return chunks

    def _get_meta_for_position(
        self,
        char_to_meta: List[Dict[str, Any]],
        position: int
    ) -> Dict[str, Any]:
        """
        Get the metadata of the block that contains the given character position.
        :param char_to_meta: List of block metadata with character ranges.
        :param position: Character position in the full text.
        :return: Metadata dict for the block at the given position.
        """
        for entry in char_to_meta:
            if entry["start"] <= position < entry["end"]:
                return entry
        # Fallback: return last block's metadata
        return char_to_meta[-1] if char_to_meta else {}